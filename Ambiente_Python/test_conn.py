import os, time, json
from pathlib import Path
import pymysql
from pymysql.cursors import DictCursor
from dotenv import load_dotenv
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =========================
# Conexão
# =========================
load_dotenv(Path(__file__).resolve().parent.parent / ".env")

def get_conn(retries=3, wait=2):
    last = None
    for _ in range(retries):
        try:
            return pymysql.connect(
                host=os.getenv("DB_HOST"),
                port=int(os.getenv("DB_PORT", "3306")),
                user=os.getenv("DB_USER"),
                password=os.getenv("DB_PASS"),
                database=os.getenv("DB_NAME"),
                charset=os.getenv("DB_CHARSET", "utf8mb4"),
                cursorclass=DictCursor,
                autocommit=False,
                connect_timeout=10, read_timeout=15, write_timeout=15,
            )
        except Exception as e:
            last = e
            time.sleep(wait)
    raise last

# =========================
# Helpers de Word
# =========================
def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_kv(doc, items):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light List Accent 1"
    for k, v in items:
        row = tbl.add_row().cells
        row[0].text = str(k)
        row[1].text = "" if v is None else str(v)
    doc.add_paragraph()

def add_table(doc, rows, cols, header=None):
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.style = "Light Grid Accent 1"
    if header:
        for j, h in enumerate(header):
            tbl.rows[0].cells[j].text = str(h)
    for r in rows:
        row = tbl.add_row().cells
        for j, val in enumerate(r):
            row[j].text = "" if val is None else str(val)
    doc.add_paragraph()

def safe(doc, title, fn):
    try:
        fn()
    except Exception as e:
        doc.add_paragraph(f"(Falha ao coletar {title}: {type(e).__name__}: {e})")

# =========================
# Coletas no information_schema
# =========================
SQL_TABLES = """
SELECT TABLE_NAME, TABLE_TYPE, ENGINE, ROW_FORMAT, TABLE_ROWS, DATA_LENGTH, INDEX_LENGTH,
       AUTO_INCREMENT, CREATE_TIME, UPDATE_TIME, TABLE_COLLATION, TABLE_COMMENT
FROM information_schema.TABLES
WHERE TABLE_SCHEMA=%s
ORDER BY TABLE_TYPE, TABLE_NAME;
"""

SQL_COLUMNS = """
SELECT COLUMN_NAME, ORDINAL_POSITION, COLUMN_TYPE, IS_NULLABLE, COLUMN_DEFAULT, EXTRA, COLUMN_KEY,
       COLLATION_NAME, COLUMN_COMMENT
FROM information_schema.COLUMNS
WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s
ORDER BY ORDINAL_POSITION;
"""

SQL_INDEXES = """
SELECT INDEX_NAME, NON_UNIQUE, SEQ_IN_INDEX, COLUMN_NAME, COLLATION, SUB_PART, INDEX_TYPE, COMMENT
FROM information_schema.STATISTICS
WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s
ORDER BY INDEX_NAME, SEQ_IN_INDEX;
"""

SQL_CONSTRAINTS = """
SELECT k.CONSTRAINT_NAME, c.CONSTRAINT_TYPE, k.COLUMN_NAME, k.REFERENCED_TABLE_NAME, k.REFERENCED_COLUMN_NAME
FROM information_schema.KEY_COLUMN_USAGE k
JOIN information_schema.TABLE_CONSTRAINTS c
  ON k.CONSTRAINT_NAME=c.CONSTRAINT_NAME AND k.TABLE_SCHEMA=c.TABLE_SCHEMA AND k.TABLE_NAME=c.TABLE_NAME
WHERE k.TABLE_SCHEMA=%s AND k.TABLE_NAME=%s
ORDER BY c.CONSTRAINT_TYPE, k.CONSTRAINT_NAME, k.ORDINAL_POSITION;
"""

SQL_VIEW_DEF = """
SELECT VIEW_DEFINITION, CHECK_OPTION, IS_UPDATABLE, DEFINER, SECURITY_TYPE
FROM information_schema.VIEWS
WHERE TABLE_SCHEMA=%s AND TABLE_NAME=%s;
"""

SQL_TRIGGERS = """
SELECT TRIGGER_NAME, EVENT_MANIPULATION, EVENT_OBJECT_TABLE, ACTION_TIMING, CREATED
FROM information_schema.TRIGGERS
WHERE TRIGGER_SCHEMA=%s
ORDER BY EVENT_OBJECT_TABLE, TRIGGER_NAME;
"""

def mermaid_er(schema_obj):
    """Gera diagrama ER (Mermaid) com FKs."""
    lines = ["```mermaid", "erDiagram"]
    # Entidades
    for t in schema_obj["tables"]:
        lines.append(f"  {t['name']} {{")
        for c in t["columns"]:
            typ = c["type"].split("(")[0].upper()
            lines.append(f"    {typ} {c['name']}")
        lines.append("  }")
    # Relacionamentos (FKs)
    for t in schema_obj["tables"]:
        for fk in (fk for fk in t["constraints"] if fk["type"] == "FOREIGN KEY" and fk["ref_table"]):
            lines.append(f"  {t['name']} }}o--|| {fk['ref_table']} : \"{fk['column']}→{fk['ref_column']}\"")
    lines.append("```")
    return "\n".join(lines)

def main():
    print("== MySQL Schema Doc Generator ==")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("SELECT NOW() AS server_time, DATABASE() AS db, VERSION() AS ver")
        meta = cur.fetchone()
        db = meta["db"]
        print(f"[INFO] Conectado em `{db}` | MySQL {meta['ver']} | server_time={meta['server_time']}")

        # =====================================================
        # Coleta objetos
        # =====================================================
        print("[STEP] Coletando objetos do information_schema...")
        cur.execute(SQL_TABLES, (db,))
        all_tables = cur.fetchall()
        base_tables = [t for t in all_tables if t["TABLE_TYPE"] == "BASE TABLE"]
        views = [t for t in all_tables if t["TABLE_TYPE"] == "VIEW"]
        print(f"[OK] Objetos: {len(base_tables)} tabelas | {len(views)} views")

        # =====================================================
        # Monta JSON de schema (para IA)
        # =====================================================
        schema = {
            "database": db,
            "mysql_version": meta["ver"],
            "generated_at": str(meta["server_time"]),
            "tables": [],
            "views": [],
            "triggers": []
        }

        # Tabelas
        print("[STEP] Lendo definição das TABELAS...")
        for i, t in enumerate(base_tables, start=1):
            name = t["TABLE_NAME"]
            print(f"  [TABELA {i}/{len(base_tables)}] {name} ...", end="", flush=True)
            # Colunas
            cur.execute(SQL_COLUMNS, (db, name))
            cols = cur.fetchall()
            # Índices
            cur.execute(SQL_INDEXES, (db, name))
            idx = cur.fetchall()
            # Constraints
            cur.execute(SQL_CONSTRAINTS, (db, name))
            cons = cur.fetchall()

            schema["tables"].append({
                "name": name,
                "engine": t["ENGINE"],
                "row_format": t["ROW_FORMAT"],
                "rows_estimated": t["TABLE_ROWS"],
                "data_length": t["DATA_LENGTH"],
                "index_length": t["INDEX_LENGTH"],
                "auto_increment": t["AUTO_INCREMENT"],
                "created": str(t["CREATE_TIME"]) if t["CREATE_TIME"] else None,
                "updated": str(t["UPDATE_TIME"]) if t["UPDATE_TIME"] else None,
                "collation": t["TABLE_COLLATION"],
                "comment": t["TABLE_COMMENT"],
                "columns": [
                    {
                        "position": c["ORDINAL_POSITION"],
                        "name": c["COLUMN_NAME"],
                        "type": c["COLUMN_TYPE"],
                        "nullable": (c["IS_NULLABLE"] == "YES"),
                        "default": c["COLUMN_DEFAULT"],
                        "extra": c["EXTRA"],
                        "key": c["COLUMN_KEY"],           # PRI/UNI/MUL
                        "collation": c["COLLATION_NAME"],
                        "comment": c["COLUMN_COMMENT"],
                    } for c in cols
                ],
                "indexes": [
                    {
                        "name": i["INDEX_NAME"],
                        "non_unique": bool(i["NON_UNIQUE"]),
                        "seq_in_index": i["SEQ_IN_INDEX"],
                        "column": i["COLUMN_NAME"],
                        "collation": i["COLLATION"],
                        "sub_part": i["SUB_PART"],
                        "index_type": i["INDEX_TYPE"],
                        "comment": i["COMMENT"],
                    } for i in idx
                ],
                "constraints": [
                    {
                        "name": c["CONSTRAINT_NAME"],
                        "type": c["CONSTRAINT_TYPE"],     # PRIMARY KEY / UNIQUE / FOREIGN KEY
                        "column": c["COLUMN_NAME"],
                        "ref_table": c["REFERENCED_TABLE_NAME"],
                        "ref_column": c["REFERENCED_COLUMN_NAME"],
                    } for c in cons
                ],
            })
            print(f" ok | {len(cols)} colunas, {len(idx)} índices, {len(cons)} constraints")

        # Views
        print("[STEP] Lendo definição das VIEWS...")
        for i, v in enumerate(views, start=1):
            vname = v["TABLE_NAME"]
            print(f"  [VIEW {i}/{len(views)}] {vname} ...", end="", flush=True)
            cur.execute(SQL_COLUMNS, (db, vname))
            vcols = cur.fetchall()
            cur.execute(SQL_VIEW_DEF, (db, vname))
            vd = cur.fetchone() or {}
            schema["views"].append({
                "name": vname,
                "columns": [
                    {
                        "position": c["ORDINAL_POSITION"],
                        "name": c["COLUMN_NAME"],
                        "type": c["COLUMN_TYPE"],
                        "nullable": (c["IS_NULLABLE"] == "YES"),
                        "default": c["COLUMN_DEFAULT"],
                    } for c in vcols
                ],
                "definition": (vd.get("VIEW_DEFINITION") or None),
                "check_option": vd.get("CHECK_OPTION"),
                "is_updatable": vd.get("IS_UPDATABLE"),
                "security_type": vd.get("SECURITY_TYPE"),
                "definer": vd.get("DEFINER"),
            })
            print(f" ok | {len(vcols)} colunas")

        # Triggers
        print("[STEP] Lendo TRIGGERS...")
        cur.execute(SQL_TRIGGERS, (db,))
        trg = cur.fetchall()
        for t in trg:
            schema["triggers"].append({
                "name": t["TRIGGER_NAME"],
                "event": t["EVENT_MANIPULATION"],
                "table": t["EVENT_OBJECT_TABLE"],
                "timing": t["ACTION_TIMING"],
                "created": str(t["CREATED"]) if t["CREATED"] else None,
            })
        print(f"[OK] Triggers: {len(schema['triggers'])} identificadas")

        # =====================================================
        # DOCX (humano)
        # =====================================================
        print("[STEP] Montando DOCX...")
        doc = Document()
        title = doc.add_heading(f"Documentação do Banco — {db}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(f"Gerado em {meta['server_time']} | MySQL {meta['ver']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        add_heading(doc, "Resumo", 1)
        add_kv(doc, [
            ("Schema", db),
            ("MySQL Version", meta["ver"]),
            ("Tabelas", len(schema["tables"])),
            ("Views", len(schema["views"])),
            ("Triggers", len(schema["triggers"])),
        ])

        add_heading(doc, "Lista de Objetos", 1)
        rows = []
        for t in all_tables:
            rows.append([
                t["TABLE_TYPE"], t["TABLE_NAME"], t["ENGINE"], t["ROW_FORMAT"],
                t["TABLE_ROWS"], t["DATA_LENGTH"], t["INDEX_LENGTH"],
                t["AUTO_INCREMENT"], t["TABLE_COLLATION"], t["TABLE_COMMENT"]
            ])
        add_table(
            doc, rows, cols=10,
            header=["Tipo","Nome","Engine","RowFmt","Rows","DataLen","IdxLen","AutoInc","Collation","Coment."]
        )

        add_heading(doc, "Detalhes de Tabelas", 1)
        for t in schema["tables"]:
            name = t["name"]
            add_heading(doc, f"Tabela: {name}", 2)
            add_kv(doc, [
                ("Engine", t["engine"]),
                ("Row Format", t["row_format"]),
                ("Rows (estimado)", t["rows_estimated"]),
                ("Data Length", t["data_length"]),
                ("Index Length", t["index_length"]),
                ("Auto Increment", t["auto_increment"]),
                ("Created", t["created"]),
                ("Updated", t["updated"]),
                ("Collation", t["collation"]),
                ("Comentário", t["comment"]),
            ])

            def add_columns():
                cols = t["columns"]
                rows = [[
                    c["position"], c["name"], c["type"],
                    "YES" if c["nullable"] else "NO",
                    c["default"], c["extra"], c["key"], c["collation"], c["comment"]
                ] for c in cols]
                add_table(doc, rows, cols=9, header=[
                    "#","Coluna","Tipo","Nulo?","Default","Extra","Key","Collation","Coment."
                ])
            safe(doc, f"colunas de {name}", add_columns)

            def add_indexes():
                idx = t["indexes"]
                if not idx:
                    doc.add_paragraph("Sem índices.")
                    return
                rows = [[
                    i["name"], int(i["non_unique"]), i["seq_in_index"], i["column"],
                    i["collation"], i["sub_part"], i["index_type"], i["comment"]
                ] for i in idx]
                add_table(doc, rows, cols=8, header=[
                    "Index","NonUnique","Seq","Coluna","Coll","SubPart","Tipo","Coment."
                ])
            safe(doc, f"índices de {name}", add_indexes)

            def add_constraints():
                cons = t["constraints"]
                if not cons:
                    doc.add_paragraph("Sem constraints declaradas.")
                    return
                rows = [[
                    c["name"], c["type"], c["column"], c["ref_table"], c["ref_column"]
                ] for c in cons]
                add_table(doc, rows, cols=5, header=[
                    "Constraint","Tipo","Coluna","Ref. Tabela","Ref. Coluna"
                ])
            safe(doc, f"constraints de {name}", add_constraints)

        add_heading(doc, "Detalhes de Views", 1)
        if not schema["views"]:
            doc.add_paragraph("Sem views no schema.")
        for v in schema["views"]:
            add_heading(doc, f"View: {v['name']}", 2)

            def add_vcols():
                rows = [[
                    c["position"], c["name"], c["type"],
                    "YES" if c["nullable"] else "NO",
                    c["default"]
                ] for c in v["columns"]]
                add_table(doc, rows, cols=5, header=["#","Coluna","Tipo","Nulo?","Default"])
            safe(doc, f"colunas de view {v['name']}", add_vcols)

            def add_vdef():
                definition = v.get("definition") or ""
                if not definition:
                    doc.add_paragraph("(Sem definição disponível ou falta permissão SHOW VIEW)")
                    return
                doc.add_paragraph("Definição (parcial):")
                if len(definition) > 4000:
                    definition = definition[:4000] + "\n... (truncado)"
                pre = doc.add_paragraph()
                run = pre.add_run(definition)
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                doc.add_paragraph()
                add_kv(doc, [
                    ("CHECK_OPTION", v.get("check_option")),
                    ("IS_UPDATABLE", v.get("is_updatable")),
                    ("SECURITY_TYPE", v.get("security_type")),
                    ("DEFINER", v.get("definer")),
                ])
            safe(doc, f"definição da view {v['name']}", add_vdef)

        add_heading(doc, "Triggers", 1)
        if not schema["triggers"]:
            doc.add_paragraph("Sem triggers no schema.")
        else:
            rows = [[
                t["name"], t["event"], t["table"], t["timing"], t["created"]
            ] for t in schema["triggers"]]
            add_table(doc, rows, cols=5, header=["Trigger","Evento","Tabela","Timing","Criada em"])

        # =====================================================
        # Salvar arquivos
        # =====================================================
        outdir = Path(__file__).resolve().parent / "documents"
        outdir.mkdir(parents=True, exist_ok=True)

        print("[STEP] Salvando saídas...")
        # 1) DOCX
        docx_path = outdir / f"mysql_documentation_{db}_{int(time.time())}.docx"
        doc.save(docx_path.as_posix())
        print(f"[OK] DOCX gerado: {docx_path}")

        # 2) JSON (para IA)
        json_path = outdir / "mysql_schema.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(schema, f, ensure_ascii=False, indent=2)
        print(f"[OK] JSON gerado: {json_path}")

        # 3) Markdown + Mermaid (ER)
        md_path = outdir / "mysql_schema.md"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# Schema: {db}\n\n")
            f.write(f"- MySQL: {meta['ver']}\n- Gerado em: {meta['server_time']}\n\n")
            f.write("## Diagrama ER (Mermaid)\n\n")
            f.write(mermaid_er(schema))
            f.write("\n\n## Resumo\n\n")
            f.write(f"- Tabelas: {len(schema['tables'])}\n- Views: {len(schema['views'])}\n- Triggers: {len(schema['triggers'])}\n")
        print(f"[OK] Markdown gerado: {md_path}")

        # 4) Dump DDL básico (schema.sql)
        ddl_path = outdir / "schema.sql"
        with open(ddl_path, "w", encoding="utf-8") as fddl:
            fddl.write(f"-- DDL dump (estrutura) de `{db}` gerado em {meta['server_time']}\n\n")
            # Tabelas
            for t in schema["tables"]:
                name = t["name"]
                try:
                    cur.execute(f"SHOW CREATE TABLE `{name}`")
                    row = cur.fetchone()
                    fddl.write(f"--\n-- Tabela `{name}`\n--\n")
                    fddl.write(row["Create Table"] + ";\n\n")
                except Exception as e:
                    fddl.write(f"-- Falha SHOW CREATE TABLE {name}: {type(e).__name__}: {e}\n\n")
            # Views
            for v in schema["views"]:
                name = v["name"]
                try:
                    cur.execute(f"SHOW CREATE VIEW `{name}`")
                    row = cur.fetchone()
                    create_view = row.get("Create View") if "Create View" in row else None
                    if create_view:
                        fddl.write(f"--\n-- View `{name}`\n--\n")
                        fddl.write(create_view + ";\n\n")
                    else:
                        if v.get("definition"):
                            fddl.write(f"--\n-- View `{name}` (definição via information_schema)\n--\n")
                            fddl.write(f"CREATE OR REPLACE VIEW `{name}` AS\n{v['definition']};\n\n")
                        else:
                            fddl.write(f"-- Sem definição disponível para view `{name}`\n\n")
                except Exception as e:
                    fddl.write(f"-- Falha SHOW CREATE VIEW {name}: {type(e).__name__}: {e}\n\n")
        print(f"[OK] DDL gerado: {ddl_path}")

        print("[DONE] Processo concluído com sucesso.")

if __name__ == "__main__":
    main()
