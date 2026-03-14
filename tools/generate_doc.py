#!/usr/bin/env python3
"""
Gera documento DOCX descritivo do OKR System.
Uso: python tools/generate_doc.py
Saida: C:/Users/wendrewgomes/Desktop/OKR_System_Documentacao.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ========== STYLES ==========
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(headers, rows, col_widths=None):
    """Helper para criar tabelas formatadas."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_db_table(name, desc, columns):
    """Helper para documentar tabela do banco."""
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    if desc:
        p.add_run(f'  —  {desc}').font.size = Pt(9)
    add_table(['Coluna', 'Tipo', 'Restricoes'], columns, [5, 4, 8])


# ============================================================
#  CAPA
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('OKR System')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PlanningBI')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0xF1, 0xC4, 0x0F)
run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentacao Tecnica do Sistema')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Estrutura, Banco de Dados, Funcionalidades e Logicas').font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(f'Gerado em: {datetime.date.today().strftime("%d/%m/%Y")}').font.size = Pt(10)

doc.add_page_break()


# ============================================================
#  SUMARIO
# ============================================================
doc.add_heading('Sumario', level=1)
toc_items = [
    '1. Visao Geral do Projeto',
    '2. Arquitetura e Tecnologias',
    '3. Estrutura de Diretorios',
    '4. Menus e Paginas do Sistema',
    '5. Banco de Dados — Tabelas, Colunas e Relacoes',
    '   5.1 Tabelas Core OKR',
    '   5.2 Milestones e Apontamentos',
    '   5.3 Iniciativas e Orcamentos',
    '   5.4 Aprovacoes',
    '   5.5 Usuarios e Autenticacao',
    '   5.6 RBAC (Controle de Acesso)',
    '   5.7 Empresa e Personalizacao',
    '   5.8 Tabelas de Dominio (Lookups)',
    '   5.9 Notificacoes',
    '   5.10 Quiz / Landing Page (LP001)',
    '6. API REST v1 — Endpoints',
    '7. Sistema RBAC (Controle de Acesso)',
    '8. Logicas de Negocio',
    '9. Deploy e Infraestrutura',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ============================================================
#  1. VISAO GERAL
# ============================================================
doc.add_heading('1. Visao Geral do Projeto', level=1)
doc.add_paragraph(
    'O OKR System e uma plataforma completa de gestao de OKRs (Objectives and Key Results) '
    'desenvolvida sob a marca PlanningBI. O sistema permite que organizacoes definam objetivos '
    'estrategicos, desdobrem em resultados-chave mensuráveis, acompanhem progresso via '
    'apontamentos e milestones, gerenciem iniciativas e orcamentos, e controlem fluxos de aprovacao.'
)
doc.add_paragraph(
    'O projeto possui tres camadas: backend web em PHP puro (sem framework), '
    'aplicativo movel em Flutter/Dart, e build web Flutter compilado. '
    'O banco de dados e MySQL hospedado na HostGator (shared hosting).'
)
doc.add_paragraph(
    'Principais funcionalidades: Dashboard analitico, Mapa Estrategico (pilares BSC), '
    'Cascata de OKRs, Gestao de Iniciativas e Tarefas, Controle Orcamentario, '
    'Sistema de Aprovacoes, Chat IA integrado (OpenAI), Relatorios PDF, '
    'Landing Page com Quiz de Maturidade OKR, e Personalizacao visual por empresa.'
)


# ============================================================
#  2. ARQUITETURA
# ============================================================
doc.add_heading('2. Arquitetura e Tecnologias', level=1)
add_table(
    ['Camada', 'Tecnologia', 'Detalhes'],
    [
        ['Backend Web', 'PHP 8.1+ (puro)', 'Sem framework; PDO/MySQL; Monolog (logging); PHPMailer (email)'],
        ['Frontend Web', 'HTML/CSS/JS + Bootstrap 5', 'Views PHP server-side; AJAX para operacoes; CSS vars para temas'],
        ['App Movel', 'Flutter 3.x / Dart', 'Comunicacao via REST API v1; ~43 arquivos .dart'],
        ['Web Build', 'Flutter Web (compilado)', 'Diretorio /app/ com build estatico'],
        ['Banco de Dados', 'MySQL 5.7+ / InnoDB', '62 tabelas + 9 views; utf8mb4_unicode_ci'],
        ['Hosting', 'HostGator (shared)', 'Deploy via git push com post-receive hook'],
        ['Autenticacao', 'Sessao PHP + JWT', 'Sessao para web; JWT (HMAC-SHA256) para API/app'],
        ['IA/Chat', 'OpenAI API', 'Chat integrado com contexto do sistema'],
    ],
    [5, 4, 9]
)


# ============================================================
#  3. ESTRUTURA DE DIRETORIOS
# ============================================================
doc.add_heading('3. Estrutura de Diretorios', level=1)
add_table(
    ['Diretorio/Arquivo', 'Descricao'],
    [
        ['/', 'Raiz do projeto = web root do servidor'],
        ['/auth/', 'Backend: config, login, registro, CRUD, ACL, mailer, helpers'],
        ['/auth/helpers/', 'Helpers de logica: cycle_calc, kr_helpers, iniciativa_envolvidos'],
        ['/api/api_platform/v1/', 'REST API v1: 59+ endpoints organizados por recurso'],
        ['/api/', 'APIs auxiliares: chat IA, avatar AI, mapa estrategico, cascata'],
        ['/views/', 'Paginas PHP do frontend web'],
        ['/views/partials/', 'Componentes compartilhados: sidebar, header, breadcrumbs, chat, tutorial'],
        ['/views/includes/', 'Helpers de view: okr_helpers'],
        ['/assets/', 'CSS, JS, imagens, avatars, company_theme.php'],
        ['/okr_app/', 'Codigo-fonte Flutter (app movel)'],
        ['/app/', 'Build compilado Flutter Web'],
        ['/LP/', 'Landing Pages (Quiz de Maturidade OKR)'],
        ['/tests/', 'Testes PHPUnit (Unit + Smoke)'],
        ['/tools/', 'Scripts auxiliares: health check, test runner'],
        ['/vendor/', 'Dependencias Composer (gitignored)'],
        ['/uploads/', 'Arquivos enviados pelos usuarios (gitignored)'],
        ['/logs/', 'Logs da aplicacao (gitignored)'],
    ],
    [5, 12]
)


# ============================================================
#  4. MENUS E PAGINAS
# ============================================================
doc.add_heading('4. Menus e Paginas do Sistema', level=1)

doc.add_heading('4.1 Estrutura do Menu (Sidebar)', level=2)
doc.add_paragraph(
    'O menu lateral e dinamico: itens sao exibidos ou ocultados conforme permissoes '
    'RBAC do usuario logado. A funcao can_open_path() consulta dom_paginas.requires_cap '
    'para decidir a visibilidade de cada item.'
)

add_table(
    ['Menu', 'Submenu', 'Arquivo', 'Visibilidade'],
    [
        ['Dashboard', '—', 'views/dashboard.php', 'ACL: R:objetivo@ORG'],
        ['Minhas Tarefas', '—', 'views/minhas_tarefas.php', 'Todos autenticados'],
        ['Mapa Estrategico', '—', 'views/mapa_estrategico.php', 'ACL: R:objetivo@ORG'],
        ['Meus OKRs', '—', 'views/meus_okrs.php', 'ACL padrao'],
        ['', 'Novo Objetivo', 'views/novo_objetivo.php', 'ACL: W:objetivo@ORG'],
        ['', 'Novo Key Result', 'views/novo_key_result.php', 'ACL: W:kr@ORG'],
        ['Orcamento', '—', 'views/orcamento.php', 'ACL padrao'],
        ['Relatorio One-Page', '—', 'views/relatorios_okrs.php', 'ACL padrao'],
        ['Aprovacoes', '—', 'views/aprovacao.php', 'ACL padrao'],
        ['Configuracoes', 'Personalizar Estilo', 'views/config_style.php', 'ACL padrao'],
        ['', 'Editar Organizacao', 'views/organizacao.php', 'ACL padrao'],
        ['', 'Gerenciar Usuarios', 'views/usuarios.php', 'ACL padrao'],
        ['Administrador', 'System Health', 'views/system_health.php', 'Somente admin_master'],
        ['', 'Empresas & Usuarios', 'views/admin_companies.php', 'Somente admin_master'],
    ],
    [4, 4, 5, 4]
)

doc.add_heading('4.2 Descricao das Paginas', level=2)

pages = [
    ('Dashboard', 'Painel principal com metricas consolidadas: total de objetivos, KRs concluidos, '
     'KRs em risco, progresso geral, distribuicao por pilar BSC. Graficos e indicadores visuais.'),
    ('Minhas Tarefas', 'Lista de tarefas do usuario corrente com filtros por status e prioridade.'),
    ('Mapa Estrategico', 'Visualizacao dos pilares BSC (Financeiro, Clientes, Processos Internos, '
     'Aprendizado e Crescimento) com objetivos conectados hierarquicamente.'),
    ('Meus OKRs / Cascata', 'Visualizacao hierarquica (cascata) dos OKRs do usuario: objetivos com '
     'seus KRs, milestones, progresso e status. Permite navegar pela arvore de OKRs.'),
    ('Novo Objetivo', 'Formulario para criacao de objetivos. Campos: descricao, tipo, pilar BSC, '
     'tipo de ciclo (anual/semestral/trimestral/bimestral/mensal/personalizado), dono, datas.'),
    ('Novo Key Result', 'Formulario de criacao de KR vinculado a um objetivo. Campos: descricao, '
     'tipo KR, natureza (binario/quantitativo/qualitativo), baseline, meta, unidade de medida, '
     'direcao da metrica, frequencia de milestones. Gera milestones automaticamente.'),
    ('Orcamento', 'Controle orcamentario por iniciativa: valor planejado, valor realizado, '
     'status financeiro, detalhamento de despesas com evidencias.'),
    ('Relatorio One-Page', 'Geracao de relatorio PDF consolidado com todos os OKRs, progresso, '
     'metricas e status de aprovacao. Suporta geracao individual e em lote.'),
    ('Aprovacoes', 'Central de aprovacoes: lista pendencias de aprovacao para objetivos, KRs, '
     'orcamentos e apontamentos. Mostra diff das alteracoes e permite aprovar/rejeitar.'),
    ('Personalizar Estilo', 'Configuracao visual da empresa: cor primaria (bg1_hex), cor secundaria '
     '(bg2_hex), upload de logo. Preview em tempo real. Aplicado via CSS variables.'),
    ('Editar Organizacao', 'Cadastro/edicao da empresa: nome fantasia, CNPJ (com validacao na '
     'Receita Federal), dados oficiais (razao social, endereco, etc.), missao e visao.'),
    ('Gerenciar Usuarios', 'CRUD de usuarios com atribuicao de roles RBAC, departamento, nivel de '
     'cargo, e override de capabilities individuais (ALLOW/DENY).'),
    ('System Health', 'Painel de saude do sistema: conexao BD, tabelas core, collation, KRs orfaos, '
     'milestones orfaos, usuarios sem role, versao PHP, espaco em disco. Execucao de testes PHPUnit.'),
    ('Empresas & Usuarios', 'Visao administrativa consolidada: KPIs (total empresas, usuarios, '
     'objetivos), lista de empresas em acordeon com usuarios vinculados, busca, badges de roles.'),
    ('Notificacoes', 'Central de notificacoes do usuario: lidas/nao lidas, com badge no header. '
     'Polling automatico a cada 60 segundos.'),
    ('Login', 'Autenticacao por email/senha com suporte a reCAPTCHA (v2 invisible, v3, Enterprise).'),
    ('Registro', 'Cadastro de novo usuario: email verificado por token, telefone/WhatsApp, empresa, '
     'avatar, senha com politica de complexidade (8+ chars, maiuscula, minuscula, numero, especial).'),
    ('Perfil do Usuario', 'Visualizacao e edicao de dados pessoais, avatar, informacoes de contato.'),
    ('Reset de Senha', 'Fluxo de recuperacao de senha com token por email, validacao e redefinicao.'),
]

for title, desc in pages:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()


# ============================================================
#  5. BANCO DE DADOS
# ============================================================
doc.add_heading('5. Banco de Dados — Tabelas, Colunas e Relacoes', level=1)
doc.add_paragraph(
    'O banco utiliza MySQL/InnoDB com charset utf8mb4_unicode_ci. Possui 62 tabelas e 9 views. '
    'As relacoes sao implementadas via Foreign Keys com acoes CASCADE e SET NULL. '
    'A seguir, a documentacao detalhada de cada tabela agrupada por dominio funcional.'
)

# --- 5.1 Core OKR ---
doc.add_heading('5.1 Tabelas Core OKR', level=2)

add_db_table('objetivos', 'Objetivos estrategicos', [
    ['id_objetivo', 'INT(10) UNSIGNED', 'PK, AUTO_INCREMENT'],
    ['descricao', 'TEXT', 'NOT NULL'],
    ['tipo', 'VARCHAR(20)', 'FK -> dom_tipo_objetivo'],
    ['pilar_bsc', 'VARCHAR(30)', 'FK -> dom_pilar_bsc'],
    ['dono', 'VARCHAR(100)', 'NOT NULL — nome do responsavel'],
    ['id_user_criador', 'INT(11)', 'FK -> usuarios'],
    ['id_company', 'INT(10) UNSIGNED', 'FK -> company'],
    ['status', 'VARCHAR(20)', 'FK -> dom_status_kr, default "Nao Iniciado"'],
    ['status_aprovacao', 'VARCHAR(15)', 'FK -> dom_status_aprovacao, default "Pendente"'],
    ['tipo_ciclo', 'VARCHAR(50)', 'FK -> dom_ciclos, NOT NULL'],
    ['ciclo', 'VARCHAR(50)', 'Descricao do ciclo'],
    ['dt_criacao', 'DATE', ''],
    ['dt_inicio', 'DATE', ''],
    ['dt_prazo', 'DATE', ''],
    ['dt_conclusao', 'DATE', ''],
    ['qualidade', 'VARCHAR(15)', 'FK -> dom_qualidade_objetivo'],
    ['observacoes', 'TEXT', ''],
    ['justificativa_ia', 'TEXT', 'Avaliacao gerada por IA'],
    ['aprovador', 'VARCHAR(100)', ''],
    ['id_user_aprovador', 'INT(11)', ''],
    ['dt_aprovacao', 'DATETIME', ''],
    ['comentarios_aprovacao', 'TEXT', ''],
])

add_db_table('key_results', 'Resultados-chave vinculados a objetivos', [
    ['id_kr', 'VARCHAR(50)', 'PK'],
    ['id_objetivo', 'INT(10) UNSIGNED', 'FK -> objetivos (CASCADE), NOT NULL'],
    ['key_result_num', 'INT(11)', 'NOT NULL — numero sequencial'],
    ['descricao', 'TEXT', 'NOT NULL'],
    ['id_user_criador', 'INT(11)', 'FK -> usuarios'],
    ['tipo_kr', 'VARCHAR(32)', 'FK -> dom_tipo_kr'],
    ['natureza_kr', 'VARCHAR(20)', 'FK -> dom_natureza_kr (binario/quantitativo/qualitativo)'],
    ['status', 'VARCHAR(20)', 'FK -> dom_status_kr'],
    ['status_aprovacao', 'VARCHAR(15)', 'FK -> dom_status_aprovacao'],
    ['tipo_frequencia_milestone', 'VARCHAR(20)', 'FK -> dom_tipo_frequencia_milestone'],
    ['baseline', 'DECIMAL(10,2)', 'Valor inicial'],
    ['meta', 'DECIMAL(10,2)', 'Valor alvo'],
    ['unidade_medida', 'VARCHAR(30)', ''],
    ['direcao_metrica', 'VARCHAR(30)', 'Crescente/Decrescente'],
    ['farol', 'VARCHAR(20)', 'Indicador visual de status'],
    ['margem_confianca', 'DECIMAL(5,2)', ''],
    ['qualidade', 'VARCHAR(15)', 'FK -> dom_qualidade_objetivo'],
    ['peso', 'DECIMAL(5,2)', 'Peso relativo no objetivo'],
    ['data_inicio', 'DATE', ''],
    ['data_fim', 'DATE', ''],
    ['responsavel', 'VARCHAR(100)', ''],
    ['dt_criacao', 'DATE', 'NOT NULL'],
    ['observacoes', 'TEXT', ''],
])

add_db_table('objetivo_links', 'Relacoes entre objetivos (cascata)', [
    ['id_link', 'BIGINT(20) UNSIGNED', 'PK, AUTO_INCREMENT'],
    ['id_company', 'INT(10) UNSIGNED', 'FK -> company, NOT NULL'],
    ['id_src', 'INT(10) UNSIGNED', 'FK -> objetivos (CASCADE) — objetivo de origem'],
    ['id_dst', 'INT(10) UNSIGNED', 'FK -> objetivos (CASCADE) — objetivo de destino'],
    ['justificativa', 'TEXT', 'NOT NULL'],
    ['ativo', 'TINYINT(1)', 'DEFAULT 1'],
    ['criado_por', 'INT(11)', ''],
    ['criado_em', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'],
])

add_db_table('okr_kr_envolvidos', 'Usuarios envolvidos em Key Results', [
    ['id', 'BIGINT(20) UNSIGNED', 'PK, AUTO_INCREMENT'],
    ['id_kr', 'VARCHAR(50)', 'FK -> key_results, NOT NULL'],
    ['id_user', 'BIGINT(20) UNSIGNED', 'FK -> usuarios, NOT NULL'],
    ['papel', 'VARCHAR(30)', 'owner|colab|viewer'],
    ['dt_incl', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'],
])

# --- 5.2 Milestones e Apontamentos ---
doc.add_heading('5.2 Milestones e Apontamentos', level=2)

add_db_table('milestones_kr', 'Marcos de acompanhamento dos KRs', [
    ['id_milestone', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['id_kr', 'VARCHAR(50)', 'FK -> key_results (CASCADE), NOT NULL'],
    ['num_ordem', 'INT(11)', 'NOT NULL — ordem sequencial'],
    ['data_ref', 'DATE', 'NOT NULL — data de referencia'],
    ['valor_esperado', 'DECIMAL(10,2)', 'NOT NULL'],
    ['valor_esperado_min', 'DECIMAL(10,2)', 'Faixa minima aceitavel'],
    ['valor_esperado_max', 'DECIMAL(10,2)', 'Faixa maxima aceitavel'],
    ['valor_real_consolidado', 'DECIMAL(10,2)', 'Valor realizado consolidado'],
    ['dt_ultimo_apontamento', 'DATETIME', ''],
    ['qtde_apontamentos', 'INT(10) UNSIGNED', 'DEFAULT 0'],
    ['gerado_automatico', 'TINYINT(1)', 'DEFAULT 1'],
    ['editado_manual', 'TINYINT(1)', 'DEFAULT 0'],
    ['bloqueado_para_edicao', 'TINYINT(1)', 'DEFAULT 0'],
    ['status_aprovacao', 'VARCHAR(15)', 'FK -> dom_status_aprovacao'],
])

add_db_table('apontamentos_kr', 'Registros de progresso (evidencias)', [
    ['id_apontamento', 'BIGINT(20) UNSIGNED', 'PK, AUTO_INCREMENT'],
    ['id_kr', 'VARCHAR(50)', 'FK -> key_results (CASCADE), NOT NULL'],
    ['id_milestone', 'INT(11)', 'FK -> milestones_kr (CASCADE)'],
    ['dt_evidencia', 'DATE', 'NOT NULL — data do fato'],
    ['dt_apontamento', 'DATETIME', 'Data do registro'],
    ['valor_real', 'DECIMAL(14,4)', 'Valor apontado'],
    ['usuario_id', 'VARCHAR(100)', 'Quem registrou'],
    ['url_evidencia', 'VARCHAR(500)', 'Link para evidencia'],
    ['observacao', 'TEXT', ''],
    ['justificativa', 'TEXT', ''],
    ['origem', 'VARCHAR(30)', 'manual|import|api|ajuste|migracao'],
])

add_db_table('apontamentos_kr_anexos', 'Arquivos anexados aos apontamentos', [
    ['id_anexo', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['id_apontamento', 'BIGINT(20) UNSIGNED', 'FK -> apontamentos_kr (CASCADE), NOT NULL'],
    ['nome_anexo', 'VARCHAR(255)', 'NOT NULL'],
    ['tipo_arquivo', 'VARCHAR(50)', 'MIME type'],
    ['caminho_arquivo', 'TEXT', 'NOT NULL — path no servidor'],
    ['tamanho_bytes', 'BIGINT(20) UNSIGNED', ''],
    ['sha256_hash', 'CHAR(64)', 'Hash de integridade'],
    ['is_deleted', 'TINYINT(1)', 'DEFAULT 0 — soft delete'],
])

# --- 5.3 Iniciativas e Orcamentos ---
doc.add_heading('5.3 Iniciativas e Orcamentos', level=2)

add_db_table('iniciativas', 'Projetos/acoes vinculados a KRs', [
    ['id_iniciativa', 'VARCHAR(100)', 'PK'],
    ['id_kr', 'VARCHAR(50)', 'FK -> key_results (CASCADE), NOT NULL'],
    ['num_iniciativa', 'INT(11)', 'NOT NULL'],
    ['descricao', 'TEXT', 'NOT NULL'],
    ['status', 'VARCHAR(20)', 'FK -> dom_status_kr'],
    ['status_aprovacao', 'VARCHAR(15)', 'FK -> dom_status_aprovacao'],
    ['id_user_criador', 'VARCHAR(50)', 'NOT NULL'],
    ['dt_criacao', 'DATE', 'NOT NULL'],
    ['dt_prazo', 'DATE', ''],
    ['id_user_responsavel', 'VARCHAR(50)', ''],
])

add_db_table('iniciativas_envolvidos', 'Participantes das iniciativas', [
    ['id_iniciativa', 'VARCHAR(100)', 'PK (composta), FK -> iniciativas (CASCADE)'],
    ['id_user', 'VARCHAR(50)', 'PK (composta)'],
    ['dt_inclusao', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'],
])

add_db_table('orcamentos', 'Orcamentos por iniciativa', [
    ['id_orcamento', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['id_iniciativa', 'VARCHAR(100)', 'NOT NULL'],
    ['valor', 'DECIMAL(12,2)', 'NOT NULL — valor planejado'],
    ['data_desembolso', 'DATE', 'NOT NULL'],
    ['status_aprovacao', 'VARCHAR(15)', 'FK -> dom_status_aprovacao'],
    ['valor_realizado', 'DECIMAL(12,2)', ''],
    ['status_financeiro', 'VARCHAR(20)', 'FK -> dom_status_financeiro'],
    ['codigo_orcamento', 'VARCHAR(100)', ''],
    ['justificativa_orcamento', 'TEXT', ''],
    ['id_user_criador', 'VARCHAR(50)', 'NOT NULL'],
    ['dt_criacao', 'DATE', 'NOT NULL'],
])

add_db_table('orcamentos_detalhes', 'Despesas detalhadas do orcamento', [
    ['id_despesa', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['id_orcamento', 'INT(11)', 'FK -> orcamentos'],
    ['valor', 'DECIMAL(10,2)', 'NOT NULL'],
    ['descricao', 'TEXT', ''],
    ['data_pagamento', 'DATE', ''],
    ['evidencia_pagamento', 'TEXT', ''],
])

# --- 5.4 Aprovacoes ---
doc.add_heading('5.4 Aprovacoes', level=2)

add_db_table('fluxo_aprovacoes', 'Fluxo de aprovacoes do sistema', [
    ['id_fluxo', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['tipo_estrutura', 'VARCHAR(30)', 'NOT NULL — objetivo/kr/orcamento/apontamento'],
    ['id_referencia', 'VARCHAR(100)', 'NOT NULL — ID do item'],
    ['tipo_operacao', 'VARCHAR(20)', 'NOT NULL — novo/alteracao'],
    ['dados_originais', 'TEXT', 'JSON com dados antes da alteracao'],
    ['dados_solicitados', 'TEXT', 'NOT NULL — JSON com novos dados'],
    ['id_user_solicitante', 'VARCHAR(50)', 'NOT NULL'],
    ['status', 'VARCHAR(15)', 'NOT NULL — pendente/aprovado/negado'],
    ['id_user_aprovador', 'VARCHAR(50)', ''],
    ['justificativa', 'TEXT', ''],
    ['data_solicitacao', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'],
    ['data_aprovacao', 'DATETIME', ''],
])

add_db_table('aprovadores', 'Usuarios aprovadores configurados', [
    ['id_aprovador', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['id_user', 'VARCHAR(50)', 'UNIQUE, NOT NULL'],
    ['tudo', 'TINYINT(1)', 'DEFAULT 0 — pode aprovar tudo'],
    ['habilitado', 'TINYINT(1)', 'DEFAULT 1'],
])

add_db_table('permissoes_aprovador', 'Tipos de aprovacao permitidos por usuario', [
    ['id_permissao', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['id_user', 'VARCHAR(50)', 'NOT NULL'],
    ['tipo_estrutura', 'VARCHAR(30)', 'NOT NULL — objetivo/kr/orcamento'],
    ['status_aprovacao', 'VARCHAR(15)', 'FK -> dom_status_aprovacao'],
])

# --- 5.5 Usuarios e Autenticacao ---
doc.add_heading('5.5 Usuarios e Autenticacao', level=2)

add_db_table('usuarios', 'Cadastro de usuarios do sistema', [
    ['id_user', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['primeiro_nome', 'VARCHAR(100)', 'NOT NULL'],
    ['ultimo_nome', 'VARCHAR(100)', ''],
    ['email_corporativo', 'VARCHAR(150)', 'UNIQUE, NOT NULL'],
    ['avatar_id', 'INT(10) UNSIGNED', 'FK -> avatars, DEFAULT 1'],
    ['telefone', 'VARCHAR(20)', ''],
    ['id_company', 'INT(10) UNSIGNED', 'FK -> company (SET NULL)'],
    ['id_departamento', 'INT(11)', 'FK -> dom_departamentos'],
    ['id_nivel_cargo', 'TINYINT(3) UNSIGNED', 'FK -> dom_niveis_cargo'],
    ['dt_cadastro', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'],
    ['id_user_criador', 'INT(11)', 'FK -> usuarios (SET NULL)'],
    ['id_permissao', 'INT(11)', 'FK -> dom_permissoes'],
])

add_db_table('usuarios_credenciais', 'Senhas (hash Argon2id)', [
    ['id_user', 'INT(11)', 'PK, FK -> usuarios (CASCADE)'],
    ['senha_hash', 'VARCHAR(255)', 'NOT NULL'],
])

add_db_table('usuarios_password_resets', 'Tokens de reset de senha', [
    ['id_reset', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['user_id', 'INT(11)', 'FK -> usuarios (CASCADE)'],
    ['token', 'VARCHAR(32)', 'NOT NULL'],
    ['selector', 'CHAR(32)', 'UNIQUE'],
    ['verifier_hash', 'CHAR(64)', ''],
    ['expira_em', 'DATETIME', 'NOT NULL'],
    ['used_at', 'DATETIME', ''],
])

add_db_table('okr_email_verifications', 'Verificacao de email no registro', [
    ['id', 'BIGINT(20) UNSIGNED', 'PK, AUTO_INCREMENT'],
    ['email', 'VARCHAR(255)', 'NOT NULL'],
    ['token', 'CHAR(64)', 'UNIQUE, NOT NULL'],
    ['code_hash', 'VARCHAR(255)', 'NOT NULL'],
    ['status', 'ENUM', 'pending|verified|expired'],
    ['attempts', 'TINYINT(3) UNSIGNED', 'DEFAULT 0'],
    ['expires_at', 'DATETIME', 'NOT NULL'],
])

# --- 5.6 RBAC ---
doc.add_heading('5.6 RBAC (Controle de Acesso Baseado em Papeis)', level=2)

add_db_table('rbac_roles', 'Papeis do sistema', [
    ['role_id', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['role_key', 'VARCHAR(50)', 'UNIQUE, NOT NULL (ex: admin_master, user_admin, gestor_master)'],
    ['role_name', 'VARCHAR(100)', 'NOT NULL — nome de exibicao'],
    ['role_desc', 'TEXT', 'Descricao do papel'],
    ['is_system', 'TINYINT(1)', 'DEFAULT 1'],
    ['is_active', 'TINYINT(1)', 'DEFAULT 1'],
])

add_db_table('rbac_capabilities', 'Capacidades/permissoes granulares', [
    ['capability_id', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['cap_key', 'VARCHAR(128)', 'UNIQUE, NOT NULL (formato: ACTION:RESOURCE@SCOPE)'],
    ['resource', 'VARCHAR(30)', 'NOT NULL — objetivo, kr, apontamento, etc.'],
    ['action', 'ENUM(R,W)', 'NOT NULL — Read ou Write'],
    ['scope', 'ENUM(OWN,TEAM,UNIT,ORG)', 'NOT NULL — escopo de abrangencia'],
    ['conditions_json', 'JSON', 'Condicoes adicionais'],
])

add_db_table('rbac_role_capability', 'Vinculo papel <-> capacidade', [
    ['role_id', 'INT(11)', 'PK (composta), FK -> rbac_roles (CASCADE)'],
    ['capability_id', 'INT(11)', 'PK (composta), FK -> rbac_capabilities (CASCADE)'],
    ['effect', 'ENUM(ALLOW,DENY)', 'DEFAULT ALLOW'],
])

add_db_table('rbac_user_role', 'Atribuicao de papel ao usuario', [
    ['user_id', 'INT(11)', 'PK (composta), FK -> usuarios (CASCADE)'],
    ['role_id', 'INT(11)', 'PK (composta), FK -> rbac_roles (CASCADE)'],
    ['valid_from', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'],
    ['valid_to', 'DATETIME', 'Vigencia (opcional)'],
])

add_db_table('rbac_user_capability', 'Override de capacidade por usuario', [
    ['user_id', 'INT(11)', 'PK (composta), FK -> usuarios (CASCADE)'],
    ['capability_id', 'INT(11)', 'PK (composta), FK -> rbac_capabilities (CASCADE)'],
    ['effect', 'ENUM(ALLOW,DENY)', 'DEFAULT ALLOW — DENY tem precedencia'],
    ['note', 'VARCHAR(255)', ''],
])

# --- 5.7 Empresa e Personalizacao ---
doc.add_heading('5.7 Empresa e Personalizacao', level=2)

add_db_table('company', 'Cadastro de organizacoes', [
    ['id_company', 'INT(10) UNSIGNED', 'PK, AUTO_INCREMENT'],
    ['organizacao', 'VARCHAR(150)', 'UNIQUE, NOT NULL — Nome Fantasia'],
    ['cnpj', 'VARCHAR(14)', 'UNIQUE (opcional)'],
    ['razao_social', 'VARCHAR(255)', 'Preenchido via consulta CNPJ'],
    ['natureza_juridica_code', 'VARCHAR(10)', ''],
    ['natureza_juridica_desc', 'VARCHAR(255)', ''],
    ['logradouro', 'VARCHAR(255)', ''],
    ['numero', 'VARCHAR(30)', ''],
    ['complemento', 'VARCHAR(100)', ''],
    ['cep', 'VARCHAR(8)', ''],
    ['bairro', 'VARCHAR(120)', ''],
    ['municipio', 'VARCHAR(120)', ''],
    ['uf', 'CHAR(2)', ''],
    ['email', 'VARCHAR(150)', ''],
    ['telefone', 'VARCHAR(30)', ''],
    ['missao', 'TEXT', ''],
    ['visao', 'TEXT', ''],
    ['situacao_cadastral', 'VARCHAR(50)', ''],
    ['data_situacao_cadastral', 'DATE', ''],
    ['created_at', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'],
    ['created_by', 'INT(10) UNSIGNED', 'NOT NULL'],
])

add_db_table('company_style', 'Personalizacao visual por empresa', [
    ['id_style', 'INT(10) UNSIGNED', 'PK, AUTO_INCREMENT'],
    ['id_company', 'INT(10) UNSIGNED', 'UNIQUE, FK -> company (CASCADE)'],
    ['bg1_hex', 'CHAR(7)', 'NOT NULL — cor primaria (escura)'],
    ['bg2_hex', 'CHAR(7)', 'NOT NULL — cor secundaria (clara/destaque)'],
    ['logo_base64', 'LONGTEXT', 'Logo em data URI (PNG/JPG/SVG)'],
    ['okr_master_user_id', 'INT(10) UNSIGNED', 'Reservado (inativo)'],
    ['created_by', 'INT(10) UNSIGNED', 'NOT NULL'],
    ['updated_by', 'INT(10) UNSIGNED', ''],
])

add_db_table('avatars', 'Avatares disponiveis para usuarios', [
    ['id', 'INT(10) UNSIGNED', 'PK, AUTO_INCREMENT'],
    ['filename', 'VARCHAR(255)', 'UNIQUE, NOT NULL'],
    ['gender', 'ENUM', 'masculino|feminino|todos'],
    ['active', 'TINYINT(1)', 'DEFAULT 1'],
])

# --- 5.8 Dominios ---
doc.add_heading('5.8 Tabelas de Dominio (Lookups)', level=2)
doc.add_paragraph(
    'Tabelas de referencia (dom_*) armazenam valores validos para campos enumerados. '
    'Sao consultadas para popular selects e validar entradas.'
)
add_table(
    ['Tabela', 'Descricao', 'Chave Primaria'],
    [
        ['dom_ciclos', 'Tipos de ciclo (anual, semestral, trimestral, bimestral, mensal, personalizado)', 'id_ciclo'],
        ['dom_pilar_bsc', 'Pilares do BSC (Financeiro, Clientes, Processos, Aprendizado)', 'id_pilar'],
        ['dom_tipo_objetivo', 'Tipos de objetivo', 'id_tipo'],
        ['dom_tipo_kr', 'Tipos de Key Result', 'id_tipo'],
        ['dom_natureza_kr', 'Natureza do KR (binario, quantitativo, qualitativo)', 'id_natureza'],
        ['dom_status_kr', 'Status de KRs/Objetivos (Nao Iniciado, Em Progresso, Concluido, Cancelado)', 'id_status'],
        ['dom_status_aprovacao', 'Status de aprovacao (pendente, aprovado, rejeitado)', 'id_status'],
        ['dom_status_financeiro', 'Status financeiro dos orcamentos', 'id_status'],
        ['dom_qualidade_objetivo', 'Qualidade do objetivo', 'id_qualidade'],
        ['dom_tipo_frequencia_milestone', 'Frequencia de milestones (semanal, quinzenal, mensal, etc.)', 'id_frequencia'],
        ['dom_modulo_aprovacao', 'Modulos que requerem aprovacao', 'id_modulo'],
        ['dom_permissoes', 'Permissoes legacy', 'id_dominio'],
        ['dom_paginas', 'Paginas do sistema com requisitos de capability', 'id_pagina'],
        ['dom_departamentos', 'Departamentos por empresa', 'id_departamento'],
        ['dom_niveis_cargo', 'Niveis hierarquicos de cargo', 'id_nivel'],
        ['dom_cargos', 'Cargos por empresa/departamento/nivel', 'id_cargo'],
    ],
    [5, 10, 3]
)

# --- 5.9 Notificacoes ---
doc.add_heading('5.9 Notificacoes', level=2)

add_db_table('notificacoes', 'Notificacoes internas do sistema', [
    ['id_notificacao', 'INT(11)', 'PK, AUTO_INCREMENT'],
    ['id_user', 'INT(11)', 'NOT NULL — destinatario'],
    ['tipo', 'VARCHAR(30)', 'NOT NULL — tipo da notificacao'],
    ['titulo', 'VARCHAR(180)', 'NOT NULL'],
    ['mensagem', 'TEXT', 'NOT NULL'],
    ['url', 'VARCHAR(255)', 'Link para acao'],
    ['lida', 'TINYINT(1)', 'DEFAULT 0'],
    ['dt_criado', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP'],
    ['dt_lida', 'DATETIME', ''],
    ['meta_json', 'JSON', 'Metadados adicionais'],
])

# --- 5.10 Quiz LP001 ---
doc.add_heading('5.10 Quiz / Landing Page (LP001)', level=2)
doc.add_paragraph(
    'Sistema completo de quiz de maturidade OKR para captacao de leads via Landing Page. '
    'Inclui versoes, dominios, perguntas, opcoes, sessoes, respostas, scoring, '
    'perfis de resultado, benchmarks, checklists, recomendacoes, A/B testing, '
    'e integracoes com email e WhatsApp.'
)

add_table(
    ['Tabela', 'Descricao'],
    [
        ['lp001_quiz', 'Quiz master (nome, slug, versao_ativa, status)'],
        ['lp001_quiz_versao', 'Versoes do quiz (id_quiz, descricao, data_publicacao)'],
        ['lp001_quiz_dominios', 'Dominios/topicos com pesos (nome, peso, ordem)'],
        ['lp001_quiz_perguntas', 'Perguntas por versao/dominio (texto, tipo, branch_key)'],
        ['lp001_quiz_opcoes', 'Opcoes de resposta (texto, explicacao, score, categoria_resposta)'],
        ['lp001_quiz_leads', 'Leads capturados (nome, email, empresa, cargo, whatsapp, UTMs)'],
        ['lp001_quiz_sessoes', 'Sessoes de quiz (lead, token, status, ip, user_agent)'],
        ['lp001_quiz_respostas', 'Respostas individuais (sessao, pergunta, opcao, score, tempo)'],
        ['lp001_quiz_scores', 'Scoring consolidado (score_total, pct, classificacao, pdf_path)'],
        ['lp001_quiz_scores_det', 'Scores detalhados por dominio'],
        ['lp001_quiz_result_profiles', 'Perfis de resultado (intervalo_score, resumo, recomendacoes)'],
        ['lp001_quiz_ab_test', 'Configuracoes de testes A/B'],
        ['lp001_quiz_benchmark_rolling', 'Benchmarks por cargo/dominio/janela temporal'],
        ['lp001_quiz_checklist_rules', 'Regras de checklist automatico (condicoes + items)'],
        ['lp001_quiz_checklist_result', 'Resultados de checklist por sessao'],
        ['lp001_quiz_recommendation_rules', 'Regras de recomendacao automatica'],
        ['lp001_quiz_agendamentos', 'Agendamentos gerados pelo quiz'],
        ['lp001_quiz_email_log', 'Log de emails enviados'],
        ['lp001_quiz_whatsapp_log', 'Log de mensagens WhatsApp'],
        ['lp001_dom_cargos', 'Cargos (lookup para LP)'],
        ['lp001_dom_faixa_colaboradores', 'Faixas de quantidade de colaboradores'],
        ['lp001_dom_faixa_faturamento', 'Faixas de faturamento'],
    ],
    [6, 11]
)


doc.add_page_break()


# ============================================================
#  6. API REST v1
# ============================================================
doc.add_heading('6. API REST v1 — Endpoints', level=1)
doc.add_paragraph(
    'A API REST v1 esta em /api/api_platform/v1/ e utiliza autenticacao JWT (HMAC-SHA256). '
    'Todas as respostas sao JSON. Suporta CORS. O roteamento e feito em index.php '
    'com resolucao de path e metodo HTTP.'
)

add_table(
    ['Recurso', 'Metodo', 'Endpoint', 'Descricao'],
    [
        ['Auth', 'POST', 'auth/login', 'Login com email/senha + reCAPTCHA'],
        ['', 'POST', 'auth/register', 'Registro de usuario'],
        ['', 'POST', 'auth/forgot-password', 'Solicitar reset de senha'],
        ['', 'POST', 'auth/reset-password', 'Completar reset de senha'],
        ['', 'GET', 'auth/me', 'Dados do usuario autenticado'],
        ['', 'PUT', 'auth/me', 'Atualizar perfil'],
        ['', 'POST/DELETE', 'auth/avatar', 'Upload/remover avatar'],
        ['Company', 'GET', 'company/me', 'Dados da empresa do usuario'],
        ['', 'PUT', 'company/me', 'Atualizar empresa'],
        ['', 'GET/PUT', 'company/style', 'Ler/atualizar estilo visual'],
        ['Objetivos', 'GET', 'objetivos', 'Listar objetivos (filtros, paginacao)'],
        ['', 'POST', 'objetivos', 'Criar objetivo'],
        ['', 'GET', 'objetivos/:id', 'Detalhe do objetivo'],
        ['', 'PUT', 'objetivos/:id', 'Atualizar objetivo'],
        ['', 'DELETE', 'objetivos/:id', 'Excluir objetivo'],
        ['KRs', 'GET', 'objetivos/:id/krs', 'Listar KRs de um objetivo'],
        ['', 'POST', 'krs', 'Criar KR (auto-gera milestones)'],
        ['', 'GET/PUT/DELETE', 'krs/:id', 'CRUD de KR individual'],
        ['', 'POST', 'krs/:id/cancelar', 'Cancelar KR'],
        ['', 'POST', 'krs/:id/reativar', 'Reativar KR'],
        ['', 'GET', 'krs/:id/milestones', 'Listar milestones do KR'],
        ['Apontamentos', 'GET', 'krs/:id/apontamentos', 'Listar apontamentos'],
        ['', 'POST', 'krs/:id/apontamentos', 'Criar apontamento(s)'],
        ['', 'DELETE', 'apontamentos/:id', 'Excluir apontamento'],
        ['Iniciativas', 'GET', 'krs/:id/iniciativas', 'Listar iniciativas do KR'],
        ['', 'POST/GET/PUT/DELETE', 'iniciativas[/:id]', 'CRUD de iniciativas'],
        ['Orcamentos', 'GET/POST', 'orcamentos', 'Listar/criar orcamentos'],
        ['', 'PUT', 'orcamentos/:id', 'Atualizar orcamento'],
        ['', 'POST', 'orcamentos/:id/despesas', 'Adicionar despesa'],
        ['Dashboard', 'GET', 'dashboard/summary', 'Metricas consolidadas'],
        ['', 'GET', 'dashboard/cascata', 'Dados da cascata'],
        ['', 'GET', 'dashboard/mapa-estrategico', 'Dados do mapa por pilar'],
        ['Aprovacoes', 'GET', 'aprovacoes', 'Listar pendencias'],
        ['', 'POST', 'aprovacoes/decidir', 'Aprovar/rejeitar'],
        ['Notificacoes', 'GET', 'notificacoes', 'Listar notificacoes'],
        ['', 'GET', 'notificacoes/count', 'Contar nao lidas'],
        ['', 'PUT', 'notificacoes/:id/lida', 'Marcar como lida'],
        ['Usuarios', 'GET/POST', 'usuarios', 'Listar/criar usuarios'],
        ['', 'GET/PUT/DELETE', 'usuarios/:id', 'CRUD usuario individual'],
        ['', 'PUT', 'usuarios/:id/role', 'Atribuir role RBAC'],
        ['Utilitarios', 'GET', 'dominios/:tabela', 'Consultar lookup tables'],
        ['', 'GET', 'responsaveis', 'Listar responsaveis'],
        ['', 'GET', 'minhas-tarefas', 'Tarefas do usuario'],
    ],
    [3, 3, 5, 7]
)


# ============================================================
#  7. SISTEMA RBAC
# ============================================================
doc.add_heading('7. Sistema RBAC (Controle de Acesso)', level=1)

doc.add_heading('7.1 Formato de Capabilities', level=2)
doc.add_paragraph(
    'As capabilities seguem o formato ACTION:RESOURCE@SCOPE. Exemplo: W:objetivo@ORG '
    'significa permissao de escrita em objetivos no escopo da organizacao.'
)
add_table(
    ['Componente', 'Valores', 'Descricao'],
    [
        ['ACTION', 'M, W, R', 'M=tudo (manage), W=escrita+leitura, R=somente leitura'],
        ['RESOURCE', 'objetivo, kr, apontamento, iniciativa, orcamento, aprovacao, user, company, relatorio, config_okrs', 'Recurso do sistema'],
        ['SCOPE', 'SYS, ORG, UNIT, TEAM, OWN', 'SYS=sistema inteiro, ORG=empresa, UNIT=departamento, TEAM=equipe, OWN=proprio'],
    ],
    [3, 6, 8]
)

doc.add_heading('7.2 Roles Existentes', level=2)
add_table(
    ['Role Key', 'Descricao', 'Comportamento'],
    [
        ['admin_master', 'Administrador geral', 'Bypass total — tem acesso a tudo sem verificacao de capabilities'],
        ['gestor_master', 'Gestor master', 'Acesso amplo, definido por capabilities vinculadas'],
        ['user_admin', 'Admin da empresa', 'Gerencia usuarios e configuracoes da empresa'],
        ['gestor_user', 'Gestor de OKRs', 'Cria/edita objetivos e KRs, aprova itens'],
        ['user_colab', 'Colaborador', 'Acesso limitado aos KRs onde esta envolvido'],
        ['user_guest', 'Convidado', 'Somente leitura'],
    ],
    [4, 4, 9]
)

doc.add_heading('7.3 Fluxo de Verificacao', level=2)
steps = [
    '1. has_cap() recebe a capability necessaria (ex: W:objetivo@ORG)',
    '2. Busca roles do usuario na tabela rbac_user_role',
    '3. Se o usuario tem role admin_master -> retorna TRUE (bypass)',
    '4. Busca capabilities de todos os roles do usuario via rbac_role_capability',
    '5. Busca overrides individuais via rbac_user_capability',
    '6. Mescla: DENY tem precedencia sobre ALLOW',
    '7. Verifica se alguma capability cobre ACTION + RESOURCE + SCOPE',
    '8. Verifica tenant (same-company) para recursos @ORG',
    '9. Regra especial: user_colab so pode W:apontamento@ORG se estiver como envolvido no KR',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')


# ============================================================
#  8. LOGICAS DE NEGOCIO
# ============================================================
doc.add_heading('8. Logicas de Negocio', level=1)

doc.add_heading('8.1 Ciclos e Datas', level=2)
doc.add_paragraph(
    'O sistema suporta 6 tipos de ciclo para objetivos: anual, semestral, trimestral, '
    'bimestral, mensal e personalizado. O helper cycle_calc.php calcula automaticamente '
    'as datas de inicio e fim com base no tipo de ciclo e na data de referencia.'
)

doc.add_heading('8.2 Geracao Automatica de Milestones', level=2)
doc.add_paragraph(
    'Ao criar um KR, o sistema gera milestones automaticamente com base na frequencia '
    'selecionada (semanal, quinzenal, mensal, etc.) e nas datas de inicio/fim do KR. '
    'Os valores esperados sao interpolados linearmente entre baseline e meta.'
)

doc.add_heading('8.3 Inferencia de Natureza do KR', level=2)
doc.add_paragraph(
    'O helper kr_helpers.php analisa a descricao e metricas do KR para inferir '
    'automaticamente a natureza: binario (sim/nao), quantitativo (numerico) ou '
    'qualitativo (escala). Isso afeta a forma de acompanhamento e visualizacao.'
)

doc.add_heading('8.4 Fluxo de Aprovacao', level=2)
doc.add_paragraph(
    'Objetivos, KRs e orcamentos passam por fluxo de aprovacao. O sistema registra '
    'o diff entre dados originais e solicitados em JSON. Aprovadores configurados '
    'podem aprovar/rejeitar com justificativa. O status_aprovacao controla o ciclo de vida.'
)

doc.add_heading('8.5 Personalizacao Visual', level=2)
doc.add_paragraph(
    'Cada empresa pode personalizar 2 cores (bg1_hex e bg2_hex) e logo. '
    'O arquivo company_theme.php gera CSS dinamico com variaveis (--bg1, --bg2, etc.) '
    'que sao consumidas por todas as paginas via link no head. '
    'O sidebar, header e componentes reagem automaticamente as cores personalizadas.'
)

doc.add_heading('8.6 Mapa Estrategico (BSC)', level=2)
doc.add_paragraph(
    'O Mapa Estrategico organiza objetivos nos 4 pilares do Balanced Scorecard: '
    'Financeiro, Clientes, Processos Internos, e Aprendizado e Crescimento. '
    'Os objetivo_links definem relacoes de causa-efeito entre objetivos de pilares diferentes.'
)

doc.add_heading('8.7 Tenant Isolation (Multi-empresa)', level=2)
doc.add_paragraph(
    'O sistema e multi-tenant: cada usuario pertence a uma empresa (id_company). '
    'Todos os recursos (objetivos, KRs, etc.) sao associados a uma empresa. '
    'O ACL verifica que o usuario so acessa recursos da propria empresa, '
    'exceto admin_master que tem visao global.'
)


# ============================================================
#  9. DEPLOY
# ============================================================
doc.add_heading('9. Deploy e Infraestrutura', level=1)

add_table(
    ['Item', 'Detalhes'],
    [
        ['Repositorio', 'GitHub: WendrewDouglas/OKR_system (origin)'],
        ['Servidor', 'HostGator shared hosting (br822.hostgator.com.br)'],
        ['Deploy remote', 'ssh://planni40@hostgator-deploy/~/repos/OKR_system.git'],
        ['Mecanismo', 'Post-receive hook: auto-checkout em ~/public_html/OKR_system/'],
        ['Comando de deploy', 'git push origin main && git push deploy main'],
        ['Web root', '/home2/planni40/public_html/OKR_system/'],
        ['PHP', '8.1+ (HostGator)'],
        ['Testes', 'PHPUnit (Unit + Smoke) via tools/run_tests.php'],
        ['Gitignored', 'vendor/, uploads/, logs/, .env, auth/outbox/'],
    ],
    [4, 13]
)


# ============================================================
#  SALVAR
# ============================================================
output_path = 'C:/Users/wendrewgomes/Desktop/OKR_System_Documentacao.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
